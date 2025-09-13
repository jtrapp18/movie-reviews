"""
Document processing utilities for extracting text and images from PDF and Word documents.
Uses temporary processing for better deployment compatibility.
"""
import os
import uuid
import tempfile
import base64
import re
from typing import Dict, Optional, Tuple, List
from werkzeug.utils import secure_filename
from docx import Document
import PyPDF2
import pdfplumber
from PIL import Image
import io


class DocumentProcessor:
    """Handles processing of uploaded documents (PDF and Word) with temporary file processing."""
    
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}
    UPLOAD_FOLDER = 'uploads/documents'
    TEMP_FOLDER = tempfile.gettempdir()
    
    @staticmethod
    def allowed_file(filename: str) -> bool:
        """Check if file extension is allowed."""
        return ('.' in filename and 
                filename.rsplit('.', 1)[1].lower() in DocumentProcessor.ALLOWED_EXTENSIONS)
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """Get file type from filename."""
        if not filename or '.' not in filename:
            return 'unknown'
        return filename.rsplit('.', 1)[1].lower()
    
    @staticmethod
    def clean_extracted_text(text: str, remove_title: bool = True) -> str:
        """
        Clean and format extracted text, optionally removing obvious titles.
        
        Args:
            text: Raw extracted text
            remove_title: Whether to attempt to remove obvious titles from the beginning
            
        Returns:
            Cleaned and formatted text
        """
        if not text:
            return ""
        
        # Normalize whitespace but preserve intentional formatting
        # Replace multiple spaces with single space, but preserve newlines and tabs
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'\n[ \t]*\n[ \t]*\n+', '\n\n', text)  # Multiple newlines to double newline
        
        # Remove obvious titles if requested
        if remove_title:
            # Split into lines for analysis
            lines = text.split('\n')
            if len(lines) > 3:  # Only process if we have enough content
                # Look for title patterns in the first few lines
                title_patterns = [
                    r'^[A-Z][A-Z\s]+$',  # ALL CAPS lines
                    r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$',  # Title Case lines
                    r'^.{1,80}$',  # Short lines (potential titles)
                ]
                
                # Check first 3 lines for title patterns
                lines_to_remove = 0
                for i, line in enumerate(lines[:3]):
                    line = line.strip()
                    if not line:  # Skip empty lines
                        lines_to_remove += 1
                        continue
                        
                    # Check if line matches title patterns
                    is_title = any(re.match(pattern, line) for pattern in title_patterns)
                    
                    # Additional heuristics for titles
                    if is_title:
                        # Check if it's followed by content (not another title)
                        next_content_line = None
                        for j in range(i + 1, min(i + 4, len(lines))):
                            if lines[j].strip():
                                next_content_line = lines[j].strip()
                                break
                        
                        # If next content line is much longer or contains common content words
                        if next_content_line and (
                            len(next_content_line) > len(line) * 2 or
                            any(word in next_content_line.lower() for word in 
                                ['the', 'and', 'of', 'to', 'in', 'for', 'with', 'by', 'is', 'are', 'was', 'were'])
                        ):
                            lines_to_remove = i + 1
                        else:
                            break
                    else:
                        break
                
                # Remove identified title lines
                if lines_to_remove > 0:
                    text = '\n'.join(lines[lines_to_remove:])
        
        # Final cleanup
        text = text.strip()
        
        # Ensure we don't have too many consecutive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text
    
    @staticmethod
    def convert_text_to_html(text: str) -> str:
        """
        Convert plain text to HTML with basic formatting preservation.
        
        Args:
            text: Plain text content
            
        Returns:
            HTML formatted text
        """
        if not text:
            return ""
        
        # Split into lines for processing
        lines = text.split('\n')
        html_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                # Empty line becomes paragraph break
                html_lines.append('<p></p>')
                continue
            
            # Check for potential headers (short lines, title case, or all caps)
            is_header = (
                len(line) < 100 and (
                    line.isupper() or  # ALL CAPS
                    re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$', line) or  # Title Case
                    re.match(r'^[A-Z][A-Z\s]+$', line)  # Mixed caps
                )
            )
            
            if is_header:
                # Determine header level based on length and case
                if line.isupper() and len(line) < 50:
                    html_lines.append(f'<h1>{line}</h1>')
                elif len(line) < 60:
                    html_lines.append(f'<h2>{line}</h2>')
                else:
                    html_lines.append(f'<h3>{line}</h3>')
            else:
                # Regular paragraph
                # Check for bold patterns (text between ** or __)
                line = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', line)
                line = re.sub(r'__(.*?)__', r'<strong>\1</strong>', line)
                
                # Check for italic patterns (text between * or _)
                line = re.sub(r'\*(.*?)\*', r'<em>\1</em>', line)
                line = re.sub(r'_(.*?)_', r'<em>\1</em>', line)
                
                html_lines.append(f'<p>{line}</p>')
        
        # Join lines and clean up
        html_content = '\n'.join(html_lines)
        
        # Remove empty paragraphs
        html_content = re.sub(r'<p></p>', '', html_content)
        
        # Clean up multiple paragraph breaks
        html_content = re.sub(r'</p>\s*<p></p>\s*<p>', '</p><p>', html_content)
        
        return html_content.strip()
    
    @staticmethod
    def generate_unique_filename(original_filename: str) -> str:
        """Generate a unique filename while preserving extension."""
        if not original_filename:
            return None
        
        # Get file extension
        file_ext = original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else ''
        
        # Generate unique name
        unique_id = str(uuid.uuid4())
        secure_name = secure_filename(original_filename.rsplit('.', 1)[0])
        
        return f"{secure_name}_{unique_id}.{file_ext}"
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file with better formatting preservation."""
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Use extract_text with layout preservation
                    page_text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
                    if page_text:
                        # Preserve the original formatting
                        text += page_text + "\n\n"  # Add double newline between pages
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF with pdfplumber: {e}")
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                    return text.strip()
            except Exception as e2:
                print(f"Error with PyPDF2 fallback: {e2}")
                return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from Word document."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    
    @staticmethod
    def extract_html_from_docx(file_path: str) -> str:
        """Extract HTML from Word document with full formatting preservation."""
        try:
            doc = Document(file_path)
            html_parts = []
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    # Empty paragraph
                    html_parts.append('<p></p>')
                    continue
                
                # Check paragraph style for headers
                style_name = paragraph.style.name.lower()
                if 'heading' in style_name or 'title' in style_name:
                    # Determine header level
                    if 'heading 1' in style_name or 'title' in style_name:
                        level = 1
                    elif 'heading 2' in style_name:
                        level = 2
                    elif 'heading 3' in style_name:
                        level = 3
                    else:
                        level = 2  # Default to h2
                    
                    # Process runs for formatting
                    header_text = DocumentProcessor._process_paragraph_runs(paragraph)
                    html_parts.append(f'<h{level}>{header_text}</h{level}>')
                else:
                    # Regular paragraph
                    para_text = DocumentProcessor._process_paragraph_runs(paragraph)
                    html_parts.append(f'<p>{para_text}</p>')
            
            return '\n'.join(html_parts)
            
        except Exception as e:
            print(f"Error extracting HTML from DOCX: {e}")
            # Fallback to plain text
            return DocumentProcessor.extract_text_from_docx(file_path)
    
    @staticmethod
    def _process_paragraph_runs(paragraph) -> str:
        """Process paragraph runs to preserve formatting."""
        html_text = ""
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Apply formatting based on run properties
            if run.bold:
                text = f'<strong>{text}</strong>'
            if run.italic:
                text = f'<em>{text}</em>'
            if run.underline:
                text = f'<u>{text}</u>'
            
            html_text += text
        
        return html_text
    
    @staticmethod
    def extract_text_from_document(file_path: str, file_type: str, clean_text: bool = True, remove_title: bool = True, convert_to_html: bool = False) -> str:
        """Extract text from document based on file type."""
        if file_type == 'pdf':
            raw_text = DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            raw_text = DocumentProcessor.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Clean the text if requested
        if clean_text and raw_text:
            cleaned_text = DocumentProcessor.clean_extracted_text(raw_text, remove_title)
        else:
            cleaned_text = raw_text
        
        # Convert to HTML if requested
        if convert_to_html and cleaned_text:
            return DocumentProcessor.convert_text_to_html(cleaned_text)
        
        return cleaned_text
    
    @staticmethod
    def extract_html_from_document(file_path: str, file_type: str, clean_text: bool = True, remove_title: bool = True) -> str:
        """Extract HTML from document with full formatting preservation."""
        if file_type == 'pdf':
            # For PDFs, extract text and convert to HTML
            raw_text = DocumentProcessor.extract_text_from_pdf(file_path)
            if clean_text and raw_text:
                cleaned_text = DocumentProcessor.clean_extracted_text(raw_text, remove_title)
            else:
                cleaned_text = raw_text
            return DocumentProcessor.convert_text_to_html(cleaned_text)
            
        elif file_type in ['docx', 'doc']:
            # For Word docs, extract directly as HTML
            return DocumentProcessor.extract_html_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def process_uploaded_document(file, upload_folder: str) -> Dict[str, str]:
        """
        Process uploaded document and return metadata.
        
        Returns:
            Dict with keys: 'filename', 'file_path', 'file_type', 'extracted_text', 'success'
        """
        if not file or not file.filename:
            return {'success': False, 'error': 'No file provided'}
        
        if not DocumentProcessor.allowed_file(file.filename):
            return {'success': False, 'error': 'File type not allowed'}
        
        # Generate unique filename
        unique_filename = DocumentProcessor.generate_unique_filename(file.filename)
        file_type = DocumentProcessor.get_file_type(file.filename)
        
        # Ensure upload directory exists
        os.makedirs(upload_folder, exist_ok=True)
        
        # Save file
        file_path = os.path.join(upload_folder, unique_filename)
        file.save(file_path)
        
        # Extract text with improved formatting and optional title removal
        try:
            extracted_text = DocumentProcessor.extract_text_from_document(file_path, file_type, clean_text=True, remove_title=True)
            if not extracted_text.strip():
                return {
                    'success': False, 
                    'error': 'Could not extract text from document',
                    'filename': unique_filename,
                    'file_path': file_path,
                    'file_type': file_type
                }
            
            return {
                'success': True,
                'filename': unique_filename,
                'file_path': file_path,
                'file_type': file_type,
                'extracted_text': extracted_text
            }
            
        except Exception as e:
            # Clean up file if text extraction failed
            if os.path.exists(file_path):
                os.remove(file_path)
            return {
                'success': False,
                'error': f'Error processing document: {str(e)}'
            }
    
    @staticmethod
    def extract_images_from_pdf(file_path: str) -> List[Dict[str, str]]:
        """Extract images from PDF and return as base64 encoded strings."""
        images = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract images from the page
                    page_images = page.images
                    for img_num, img in enumerate(page_images):
                        try:
                            # Get image data
                            img_data = page.crop(img).to_image()
                            if img_data:
                                # Convert to PIL Image
                                pil_img = img_data.original
                                
                                # Convert to base64
                                buffer = io.BytesIO()
                                pil_img.save(buffer, format='PNG')
                                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                                
                                images.append({
                                    'page': page_num + 1,
                                    'image_num': img_num + 1,
                                    'data': f"data:image/png;base64,{img_base64}",
                                    'width': pil_img.width,
                                    'height': pil_img.height
                                })
                        except Exception as e:
                            print(f"Error extracting image {img_num} from page {page_num}: {e}")
                            continue
        except Exception as e:
            print(f"Error extracting images from PDF: {e}")
        
        return images
    
    @staticmethod
    def extract_images_from_docx(file_path: str) -> List[Dict[str, str]]:
        """Extract images from Word document and return as base64 encoded strings."""
        images = []
        try:
            doc = Document(file_path)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref.content_type:
                    try:
                        image_data = rel.target_part.blob
                        img_base64 = base64.b64encode(image_data).decode('utf-8')
                        
                        # Determine image format
                        img_format = 'png'
                        if rel.target_ref.content_type == 'image/jpeg':
                            img_format = 'jpeg'
                        elif rel.target_ref.content_type == 'image/png':
                            img_format = 'png'
                        elif rel.target_ref.content_type == 'image/gif':
                            img_format = 'gif'
                        
                        images.append({
                            'image_num': len(images) + 1,
                            'data': f"data:image/{img_format};base64,{img_base64}",
                            'format': img_format
                        })
                    except Exception as e:
                        print(f"Error processing image in Word doc: {e}")
                        continue
        except Exception as e:
            print(f"Error extracting images from Word document: {e}")
        
        return images
    
    @staticmethod
    def extract_images_from_document(file_path: str, file_type: str) -> List[Dict[str, str]]:
        """Extract images from document based on file type."""
        if file_type == 'pdf':
            return DocumentProcessor.extract_images_from_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return DocumentProcessor.extract_images_from_docx(file_path)
        else:
            return []
    
    @staticmethod
    def process_uploaded_document_temporary(file) -> Dict[str, any]:
        """
        Process uploaded document using temporary files for better deployment compatibility.
        Extracts text and keeps file for serving - deployment-friendly approach.
        
        Returns:
            Dict with keys: 'filename', 'file_path', 'file_type', 'extracted_text', 'success'
        """
        if not file or not file.filename:
            return {'success': False, 'error': 'No file provided'}
        
        if not DocumentProcessor.allowed_file(file.filename):
            return {'success': False, 'error': 'File type not allowed'}
        
        file_type = DocumentProcessor.get_file_type(file.filename)
        original_filename = secure_filename(file.filename)
        
        # Create temporary file
        temp_file = None
        try:
            # Create temporary file with proper extension
            temp_fd, temp_path = tempfile.mkstemp(suffix=f'.{file_type}')
            temp_file = os.fdopen(temp_fd, 'wb')
            
            # Save uploaded file to temporary location
            file.seek(0)  # Reset file pointer
            temp_file.write(file.read())
            temp_file.close()
            
            # Extract text with improved formatting and optional title removal
            extracted_text = DocumentProcessor.extract_text_from_document(temp_path, file_type, clean_text=True, remove_title=True)
            if not extracted_text.strip():
                return {
                    'success': False, 
                    'error': 'Could not extract text from document',
                    'filename': original_filename,
                    'file_type': file_type
                }
            
            return {
                'success': True,
                'filename': original_filename,
                'file_path': temp_path,  # Return the temp file path for serving
                'file_type': file_type,
                'extracted_text': extracted_text
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing document: {str(e)}'
            }
        finally:
            # Don't clean up the temporary file - we need it for serving
            if temp_file:
                temp_file.close()
    
    @staticmethod
    def get_document_preview(file_path: str, file_type: str, max_chars: int = 500) -> str:
        """Get a preview of the document content."""
        try:
            full_text = DocumentProcessor.extract_text_from_document(file_path, file_type, clean_text=True, remove_title=True)
            if len(full_text) <= max_chars:
                return full_text
            return full_text[:max_chars] + "..."
        except Exception as e:
            return f"Error generating preview: {str(e)}"
